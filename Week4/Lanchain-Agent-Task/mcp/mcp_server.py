"""
MCP Server for Google Docs Integration
Provides HTTP endpoints to search and read documents from Google Drive folder
"""
from http.server import HTTPServer, BaseHTTPRequestHandler
import json
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config_loader

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
from docx import Document
import PyPDF2

SCOPES = ["https://www.googleapis.com/auth/drive.readonly", 
          "https://www.googleapis.com/auth/documents.readonly"]

# Initialize Google API services
try:
    credentials = service_account.Credentials.from_service_account_file(
        "service_account.json",
        scopes=SCOPES
    )
    drive_service = build("drive", "v3", credentials=credentials)
    docs_service = build("docs", "v1", credentials=credentials)
    print("-> Google API services initialized")
except Exception as e:
    print(f"Warning: Could not initialize Google API services: {e}")
    drive_service = None
    docs_service = None


class MCPRequestHandler(BaseHTTPRequestHandler):
    """HTTP Request Handler for MCP Server"""
    
    def _send_json_response(self, status_code, data):
        """Send JSON response to client"""
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(data, indent=2).encode())
    
    def _extract_text_from_doc(self, document):
        """Extract plain text from Google Docs document structure"""
        content = document.get('body', {}).get('content', [])
        text_parts = []
        
        for element in content:
            if 'paragraph' in element:
                for text_run in element['paragraph'].get('elements', []):
                    if 'textRun' in text_run:
                        text_parts.append(text_run['textRun'].get('content', ''))
        
        return ''.join(text_parts)
    
    def _extract_text_from_docx(self, file_id):
        """Extract plain text from Word document (.docx) by downloading it"""
        try:
            # Download the file
            request = drive_service.files().get_media(fileId=file_id)
            file_buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(file_buffer, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            # Read the Word document
            file_buffer.seek(0)
            doc = Document(file_buffer)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        
        except Exception as e:
            print(f"  -> Error extracting text from Word document: {e}")
            return ""
    
    def _extract_text_from_pdf(self, file_id):
        """Extract plain text from PDF document by downloading it"""
        try:
            # Download the file
            request = drive_service.files().get_media(fileId=file_id)
            file_buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(file_buffer, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            # Read the PDF
            file_buffer.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_buffer)
            
            # Extract text from all pages
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            return '\n'.join(text_parts)
        
        except Exception as e:
            print(f"  -> Error extracting text from PDF: {e}")
            return ""
    
    def _search_folder_documents(self, folder_name, query):
        """Search documents in a specific Google Drive folder"""
        if not drive_service or not docs_service:
            return {
                'error': 'Google API services not initialized',
                'details': 'Check service_account.json file'
            }
        
        try:
            # Find folder by name
            print(f"-> Searching for folder: '{folder_name}'")
            folder_results = drive_service.files().list(
                q=f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'",
                fields="files(id, name)"
            ).execute()
            
            folders = folder_results.get('files', [])
            print(f"-> Found {len(folders)} folder(s) matching '{folder_name}'")
            
            if not folders:
                # List all accessible folders for debugging
                print("-> Listing all accessible folders:")
                all_folders = drive_service.files().list(
                    q="mimeType='application/vnd.google-apps.folder'",
                    fields="files(id, name)",
                    pageSize=20
                ).execute()
                for f in all_folders.get('files', []):
                    print(f"  - {f['name']} (ID: {f['id']})")
                
                return {
                    'error': f"Folder '{folder_name}' not found. Check that: 1) Folder exists, 2) Folder is shared with service account, 3) Folder name matches exactly",
                    'query': query,
                    'accessible_folders': [f['name'] for f in all_folders.get('files', [])]
                }
            
            folder_id = folders[0]['id']
            print(f"-> Using folder: '{folders[0]['name']}' (ID: {folder_id})")
            
            # First, list ALL files in the folder for debugging
            print("-> Listing all files in folder...")
            all_files_results = drive_service.files().list(
                q=f"'{folder_id}' in parents",
                fields="files(id, name, mimeType, createdTime, modifiedTime)"
            ).execute()
            
            all_files = all_files_results.get('files', [])
            print(f"-> Found {len(all_files)} total file(s) in folder:")
            for f in all_files:
                print(f"  - {f['name']} (Type: {f['mimeType']})")
            
            # Search for documents in the folder (Google Docs, Word, PDF)
            print("-> Searching for documents (Google Docs, Word, PDF)...")
            docs_query = (
                f"'{folder_id}' in parents and "
                f"(mimeType='application/vnd.google-apps.document' or "
                f"mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                f"mimeType='application/pdf')"
            )
            docs_results = drive_service.files().list(
                q=docs_query,
                fields="files(id, name, mimeType, createdTime, modifiedTime)"
            ).execute()
            
            docs = docs_results.get('files', [])
            print(f"-> Found {len(docs)} document(s)")
            for doc in docs:
                print(f"  - {doc['name']} ({doc['mimeType']})")
            
            if not docs:
                return {
                    'message': f"No documents found in folder '{folder_name}'",
                    'query': query,
                    'documents': []
                }
            
            # Read content from each document and search
            matched_docs = []
            all_docs_content = []  # Store all docs for fallback
            query_lower = query.lower()
            
            print(f"-> Extracting content from {len(docs)} document(s)...")
            
            for doc in docs:
                doc_id = doc['id']
                doc_mime_type = doc['mimeType']
                doc_name = doc['name']
                
                print(f"  -> Processing: {doc_name}")
                
                # Extract content based on document type
                if doc_mime_type == 'application/vnd.google-apps.document':
                    # Google Docs
                    document = docs_service.documents().get(documentId=doc_id).execute()
                    content = self._extract_text_from_doc(document)
                elif doc_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    # Word document
                    print("Downloading Word document...")
                    content = self._extract_text_from_docx(doc_id)
                elif doc_mime_type == 'application/pdf':
                    # PDF document
                    print("Downloading PDF document...")
                    content = self._extract_text_from_pdf(doc_id)
                else:
                    print(f"  -> Skipping unsupported type: {doc_mime_type}")
                    continue
                
                print(f"     Extracted {len(content)} characters")
                print(f"     First 200 chars: {content[:200]}")
                
                # Store for potential fallback
                doc_info = {
                    'id': doc_id,
                    'title': doc_name,
                    'content': content[:2000] + ('...' if len(content) > 2000 else ''),
                    'full_content': content,
                    'created': doc.get('createdTime', ''),
                    'modified': doc.get('modifiedTime', '')
                }
                all_docs_content.append(doc_info)
                
                # Check if query matches (flexible keyword search)
                # Split query into keywords and check if ANY keyword matches
                query_keywords = query_lower.split()
                content_lower = content.lower()
                name_lower = doc_name.lower()
                
                # Match if any keyword is in content or filename
                has_match = any(keyword in content_lower or keyword in name_lower 
                               for keyword in query_keywords)
                
                if has_match:
                    print(f"     ✓ MATCH found for query: {query}")
                    matched_docs.append(doc_info)
                else:
                    print(f"     ✗ No match for query keywords: {query_keywords}")
            
            print(f"-> Search complete: {len(matched_docs)} matches out of {len(all_docs_content)} documents")
            
            # If no matches found, return all documents (relaxed search)
            if len(matched_docs) == 0 and len(all_docs_content) > 0:
                print("-> No keyword matches, returning all documents in folder")
                matched_docs = all_docs_content
            
            return {
                'query': query,
                'folder': folder_name,
                'total_results': len(matched_docs),
                'documents': matched_docs
            }
        
        except Exception as e:
            return {
                'error': f'Error searching documents: {str(e)}',
                'query': query
            }
    
    def do_POST(self):
        """Handle POST requests"""
        if self.path == '/invoke':
            try:
                # Read request body
                content_length = int(self.headers['Content-Length'])
                post_data = self.rfile.read(content_length)
                request_data = json.loads(post_data.decode('utf-8'))
                
                tool_name = request_data.get('tool')
                arguments = request_data.get('arguments', {})
                
                # Handle Google_Docs_Search tool
                if tool_name == 'Google_Docs_Search':
                    query = arguments.get('query', '')
                    folder_name = config_loader.GOOGLE_DOCS_FOLDER_NAME
                    
                    result = self._search_folder_documents(folder_name, query)
                    self._send_json_response(200, result)
                else:
                    self._send_json_response(400, {
                        'error': f'Unknown tool: {tool_name}'
                    })
            
            except Exception as e:
                self._send_json_response(500, {
                    'error': f'Server error: {str(e)}'
                })
        else:
            self._send_json_response(404, {
                'error': 'Endpoint not found',
                'available_endpoints': ['POST /invoke']
            })
    
    def do_GET(self):
        """Handle GET requests"""
        if self.path == '/health':
            self._send_json_response(200, {
                'status': 'healthy',
                'service': 'MCP Google Docs Server',
                'google_api_ready': drive_service is not None and docs_service is not None
            })
        else:
            self._send_json_response(404, {
                'error': 'Use POST /invoke to call tools'
            })
    
    def log_message(self, format, *args):
        """Custom log format"""
        print(f"[MCP Server] {format % args}")


def start_server():
    """Start MCP Server"""
    config_loader.load_config()
    host = config_loader.MCP_SERVER_HOST
    port = config_loader.MCP_SERVER_PORT
    
    server_address = (host, port)
    httpd = HTTPServer(server_address, MCPRequestHandler)
    
    print("-> MCP SERVER - Google Docs Integration")
    print(f"Server: http://{host}:{port}")
    print(f"Health: http://{host}:{port}/health")
    print(f"Folder: {config_loader.GOOGLE_DOCS_FOLDER_NAME}")
    print("\nPress Ctrl+C to stop\n")
    
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\n\nShutting down server...")
        httpd.shutdown()
        print("Server stopped\n")


if __name__ == "__main__":
    start_server()
